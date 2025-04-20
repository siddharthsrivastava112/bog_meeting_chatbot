import os
import re
import nltk
import wordninja
import fitz  # PyMuPDF for PDF parsing
from tqdm import tqdm
from docx import Document as DocxDocument
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter

# Ensure sentence tokenizer is available
nltk.download('punkt')
from nltk.tokenize import sent_tokenize

# Configuration
BASE_VECTOR_STORE = "vector_store"
DATA_FOLDER = "data"
EMBEDDING_MODEL = "sentence-transformers/all-MiniLM-L6-v2"  # 🔁 Changed model
CHUNK_SIZE = 1600  # ⬆ Increased chunk size
CHUNK_OVERLAP = 200
MIN_CHUNK_LENGTH = 300
MAX_CHUNK_LENGTH = 2000
COMBINED_INDEX_NAME = "db_faiss"

def clean_text(text):
    if not text:
        return ""
    text = re.sub(r'[\x00-\x1F\x7F-\x9F]', '', text)
    text = re.sub(r'(?<=\w)-\s+(?=\w)', '', text)
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[•▪�]+', ' ', text)
    text = re.sub(r'(?i)\b(?:page|pase)\s*\d+\b', '', text)
    text = re.sub(r'_{2,}', '', text)
    return text.strip()

def fix_spacing_in_chunk(chunk):
    if " " not in chunk:
        return " ".join(wordninja.split(chunk))
    words = chunk.split()
    fixed_words = []
    for word in words:
        if len(word) > 15 and word.isalpha():
            split = wordninja.split(word)
            fixed_words.extend(split if len(split) > 1 else [word])
        else:
            fixed_words.append(word)
    return " ".join(fixed_words)

def detect_item_and_resolution(text):
    item_match = re.search(r'Item No\.\s*\d+', text, re.IGNORECASE)
    resolution_match = re.search(r'Resolution with respect to[:\-]?\s*(.+)', text, re.IGNORECASE)
    return (
        item_match.group(0) if item_match else None,
        resolution_match.group(1).strip() if resolution_match else None
    )

def split_text_into_chunks(text, metadata):
    sentences = sent_tokenize(text)
    chunks = []
    current_chunk = ""

    for sentence in sentences:
        sentence = sentence.strip()
        if not sentence:
            continue

        if len(current_chunk) + len(sentence) + 1 <= CHUNK_SIZE:
            current_chunk += " " + sentence
        else:
            fixed_chunk = fix_spacing_in_chunk(current_chunk.strip())
            if MIN_CHUNK_LENGTH <= len(fixed_chunk) <= MAX_CHUNK_LENGTH:
                chunks.append(Document(page_content=fixed_chunk, metadata=metadata))
            current_chunk = sentence

    if current_chunk.strip():
        fixed_chunk = fix_spacing_in_chunk(current_chunk.strip())
        if MIN_CHUNK_LENGTH <= len(fixed_chunk) <= MAX_CHUNK_LENGTH:
            chunks.append(Document(page_content=fixed_chunk, metadata=metadata))

    return chunks

def process_docx_paragraphs(docx_path):
    doc = DocxDocument(docx_path)
    file_name = os.path.basename(docx_path)

    all_chunks = []
    buffer = ""
    buffer_meta = {"source": file_name}
    total_length = 0

    for para in doc.paragraphs:
        cleaned = clean_text(para.text)
        if not cleaned:
            continue

        item, resolution = detect_item_and_resolution(cleaned)
        meta = {
            "source": file_name,
            "item_no": item,
            "resolution": resolution
        }

        if len(buffer) + len(cleaned) < CHUNK_SIZE:
            buffer += " " + cleaned
            total_length += len(cleaned)
            if item:
                buffer_meta["item_no"] = item
            if resolution:
                buffer_meta["resolution"] = resolution
        else:
            all_chunks.extend(split_text_into_chunks(buffer.strip(), buffer_meta))
            buffer = cleaned
            buffer_meta = meta
            total_length = len(cleaned)

    if buffer.strip():
        all_chunks.extend(split_text_into_chunks(buffer.strip(), buffer_meta))

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(clean_text(cell.text) for cell in row.cells if clean_text(cell.text))
            if row_text:
                meta = {"source": file_name}
                all_chunks.extend(split_text_into_chunks(row_text, meta))

    return all_chunks

def process_pdf_text(pdf_path):
    doc = fitz.open(pdf_path)
    file_name = os.path.basename(pdf_path)
    all_chunks = []
    full_text = ""

    for page in doc:
        text = clean_text(page.get_text("text"))
        full_text += " " + text

    meta = {"source": file_name}
    all_chunks.extend(split_text_into_chunks(full_text.strip(), meta))
    return all_chunks

def store_per_doc_embeddings(embedder):
    files = [f for f in os.listdir(DATA_FOLDER) if f.lower().endswith(('.docx', '.pdf'))]
    all_chunks = []

    if not files:
        print("No supported files found.")
        return

    for file in tqdm(files, desc="Processing Files"):
        file_path = os.path.join(DATA_FOLDER, file)

        if file.lower().endswith('.docx'):
            chunks = process_docx_paragraphs(file_path)
        elif file.lower().endswith('.pdf'):
            chunks = process_pdf_text(file_path)
        else:
            continue

        if not chunks:
            print(f"⚠️ No valid chunks for {file}")
            continue

        all_chunks.extend(chunks)

        try:
            vector_store = FAISS.from_documents(chunks, embedder)
            doc_vector_dir = os.path.join(BASE_VECTOR_STORE, os.path.splitext(file)[0])
            os.makedirs(doc_vector_dir, exist_ok=True)
            vector_store.save_local(doc_vector_dir)
            print(f"✅ Indexed {len(chunks)} chunks from {file}")
        except Exception as e:
            print(f"❌ Failed to embed/index {file}: {e}")

    return all_chunks

def store_combined_embeddings(all_chunks, embedder):
    if not all_chunks:
        print("❌ No chunks to combine.")
        return
    try:
        combined_index = FAISS.from_documents(all_chunks, embedder)
        combined_index.save_local(os.path.join(BASE_VECTOR_STORE, COMBINED_INDEX_NAME))
        print(f"✅ Combined index saved to '{COMBINED_INDEX_NAME}'")
    except Exception as e:
        print(f"❌ Failed to create combined FAISS index: {e}")

if __name__ == "__main__":
    embedder = HuggingFaceEmbeddings(
        model_name=EMBEDDING_MODEL,
        encode_kwargs={"normalize_embeddings": True}
    )

    all_chunks = store_per_doc_embeddings(embedder)
    store_combined_embeddings(all_chunks, embedder)
